# SPDX-License-Identifier: Apache-2.0
#
# SPDX-FileCopyrightText: © 2024 Tenstorrent AI ULC

import os
from typing import List, Optional, Dict, Any
import pypdf
import docx
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from langchain_text_splitters import (
    HTMLHeaderTextSplitter,
    Language,
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)
from shared_config.logger_config import get_logger

logger = get_logger(__name__)

# Heading levels that define section boundaries for markdown/HTML splitting.
_MD_HEADERS = [("#", "h1"), ("##", "h2"), ("###", "h3")]
_HTML_HEADERS = [("h1", "h1"), ("h2", "h2"), ("h3", "h3")]

# Language-aware splitting for code files so functions/classes stay intact.
LANGUAGE_BY_EXTENSION = {
    '.py': Language.PYTHON,
    '.js': Language.JS,
    '.jsx': Language.JS,
    '.ts': Language.TS,
    '.tsx': Language.TS,
}


def _collapse_section(meta: Dict[str, Any]) -> Optional[str]:
    """Join h1/h2/h3 splitter keys into one "H1 > H2 > H3" section string."""
    parts = [meta.get(level) for level in ("h1", "h2", "h3")]
    section = " > ".join(p for p in parts if p)
    return section or None


class DocumentProcessor:
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'application/pdf',
        '.txt': 'text/plain',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.md': 'text/markdown',
        '.html': 'text/html',
        '.py': 'text/x-python',
        '.js': 'application/javascript',
        '.ts': 'application/typescript',
        '.tsx': 'application/typescript',
        '.jsx': 'application/javascript',
    }

    @staticmethod
    def get_file_type(file_path: str) -> Optional[str]:
        """Detect file type based on extension and mime type."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext in DocumentProcessor.SUPPORTED_EXTENSIONS:
            return DocumentProcessor.SUPPORTED_EXTENSIONS[ext]
        return None

    @staticmethod
    def process_pdf(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process PDF files."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                return DocumentProcessor.pages_to_documents(
                    (page.extract_text() for page in pdf_reader.pages), metadata
                )
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise

    @staticmethod
    def pages_to_documents(page_texts, metadata: Dict[str, Any]) -> List[Document]:
        """One Document per non-empty page, each with its own metadata + page number."""
        documents = []
        for page_number, text in enumerate(page_texts, start=1):
            if not text or not text.strip():
                continue
            documents.append(
                Document(page_content=text, metadata={**metadata, "page": page_number})
            )
        return documents

    @staticmethod
    def process_text(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                return [Document(page_content=content, metadata=metadata)]
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
                return [Document(page_content=content, metadata=metadata)]

    @staticmethod
    def process_docx(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process Word documents."""
        try:
            doc = docx.Document(file_path)
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise

    @staticmethod
    def process_markdown(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process Markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return DocumentProcessor.split_markdown_sections(content, metadata)
        except Exception as e:
            logger.error(f"Error processing Markdown file {file_path}: {str(e)}")
            raise

    @staticmethod
    def split_markdown_sections(content: str, metadata: Dict[str, Any]) -> List[Document]:
        """Split markdown into heading-aligned sections carrying `section` metadata."""
        splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=_MD_HEADERS, strip_headers=False
        )
        documents = []
        for piece in splitter.split_text(content):
            section = _collapse_section(piece.metadata)
            piece_metadata = dict(metadata)
            if section:
                piece_metadata["section"] = section
            documents.append(
                Document(page_content=piece.page_content, metadata=piece_metadata)
            )
        return documents or [Document(page_content=content, metadata=metadata)]

    @staticmethod
    def process_html(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process HTML files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            soup = BeautifulSoup(content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            try:
                return DocumentProcessor._split_html_sections(str(soup), metadata)
            except Exception as e:
                logger.warning(
                    f"Heading-aware HTML split failed for {file_path} ({e}); "
                    "falling back to plain text"
                )
                text = soup.get_text(separator='\n')
                return [Document(page_content=text, metadata=metadata)]
        except Exception as e:
            logger.error(f"Error processing HTML file {file_path}: {str(e)}")
            raise

    @staticmethod
    def _split_html_sections(html: str, metadata: Dict[str, Any]) -> List[Document]:
        """Split HTML on h1-h3 boundaries; the splitter emits one piece per element,
        so consecutive pieces under the same headings are re-joined into sections."""
        pieces = HTMLHeaderTextSplitter(_HTML_HEADERS).split_text(html)
        documents = []
        current_section = None
        current_texts: List[str] = []

        def flush():
            if not current_texts:
                return
            piece_metadata = dict(metadata)
            if current_section:
                piece_metadata["section"] = current_section
            documents.append(
                Document(page_content="\n".join(current_texts), metadata=piece_metadata)
            )

        for piece in pieces:
            section = _collapse_section(piece.metadata)
            if section != current_section:
                flush()
                current_section = section
                current_texts = []
            if piece.page_content.strip():
                current_texts.append(piece.page_content)
        flush()
        if not documents:
            raise ValueError("HTML splitter produced no content")
        return documents

    @staticmethod
    def process_code(file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process code files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            logger.error(f"Error processing code file {file_path}: {str(e)}")
            raise

    @classmethod
    def process_document(cls, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process document based on file type."""
        file_type = cls.get_file_type(file_path)
        if not file_type:
            raise ValueError(f"Unsupported file type: {file_path}")

        processors = {
            'application/pdf': cls.process_pdf,
            'text/plain': cls.process_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': cls.process_docx,
            'text/markdown': cls.process_markdown,
            'text/html': cls.process_html,
        }

        # Handle code files (.py, .js, .jsx, .ts, .tsx)
        code_mime_types = {'text/x-python', 'application/javascript', 'application/typescript'}
        if file_type in code_mime_types:
            return cls.process_code(file_path, metadata)

        if file_type not in processors:
            raise ValueError(f"No processor available for file type: {file_type}")

        return processors[file_type](file_path, metadata)

    @staticmethod
    def chunk_documents(
        documents: List[Document],
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        language: Optional[Language] = None,
    ) -> List[Document]:
        """Split documents into chunks, language-aware when splitting code."""
        if language is not None:
            text_splitter = RecursiveCharacterTextSplitter.from_language(
                language, chunk_size=chunk_size, chunk_overlap=chunk_overlap
            )
        else:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size, chunk_overlap=chunk_overlap
            )
        return text_splitter.split_documents(documents)